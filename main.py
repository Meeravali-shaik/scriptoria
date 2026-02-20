"""Coffee-with-Cinema: AI Storyboard & Script Generator — Flask Backend

Implements:
    1) Document generation (TXT / PDF / DOCX)
    2) RESTful communication layer (JSON APIs)
    3) Session-based state management

Endpoints:
    - GET  /                 Health check (JSON)
    - POST /set_username     Store username in session
    - POST /generate_content Store generated content in session (no AI logic here)
    - POST /download/<fmt>   Download stored outputs as txt/pdf/docx

Notes:
    - Uses server-side sessions (filesystem) to safely store large generated outputs.
    - Secret key must be provided via environment variable: FLASK_SECRET_KEY
    - CORS is configurable via environment variable: CORS_ORIGINS
"""

from __future__ import annotations

import io
import os
import secrets
from datetime import datetime, timezone
from typing import Any, Dict, Tuple

from flask import Flask, jsonify, render_template, request, session, send_file

from flask_cors import CORS
from flask_session import Session

from dotenv import load_dotenv


# -----------------------------
# Flask app factory
# -----------------------------

def create_app() -> Flask:
    # Load local environment variables from .env (safe no-op if missing).
    load_dotenv(override=False)

    app = Flask(__name__)

    debug_flag = os.getenv("FLASK_DEBUG", "").strip().lower() in {"1", "true", "yes", "on"}
    app.config["DEBUG"] = debug_flag

    secret = os.getenv("FLASK_SECRET_KEY", "").strip()
    if not secret:
        if debug_flag:
            secret = secrets.token_urlsafe(32)
            app.logger.warning(
                "FLASK_SECRET_KEY is not set; using a temporary in-memory secret because FLASK_DEBUG is enabled. "
                "Set FLASK_SECRET_KEY for stable sessions."
            )
        else:
            raise RuntimeError("Missing FLASK_SECRET_KEY environment variable.")

    app.config.update(
        SECRET_KEY=secret,
        JSON_SORT_KEYS=False,
        SESSION_COOKIE_HTTPONLY=True,
        SESSION_COOKIE_SAMESITE="Lax",
    )

    # Server-side sessions (filesystem) to avoid cookie-size limits.
    session_dir = os.path.join(app.instance_path, "flask_session")
    os.makedirs(session_dir, exist_ok=True)
    app.config.update(
        SESSION_TYPE="filesystem",
        SESSION_FILE_DIR=session_dir,
        SESSION_PERMANENT=False,
        SESSION_USE_SIGNER=True,
    )
    Session(app)

    # Cookie security can be toggled for HTTPS deployments.
    cookie_secure = os.getenv("FLASK_COOKIE_SECURE", "").strip().lower() in {"1", "true", "yes", "on"}
    if cookie_secure:
        app.config["SESSION_COOKIE_SECURE"] = True

    # CORS (optional). Only enable if CORS_ORIGINS is provided.
    # For same-origin usage (serving the frontend from this Flask app), CORS is not needed.
    origins_raw = os.getenv("CORS_ORIGINS", "").strip()
    if origins_raw:
        origins = [o.strip() for o in origins_raw.split(",") if o.strip()]
        CORS(app, resources={r"/*": {"origins": origins}}, supports_credentials=True)

    register_routes(app)
    register_error_handlers(app)
    return app


# -----------------------------
# Helpers
# -----------------------------

def _json_error(message: str, *, status: int = 400, code: str = "bad_request"):
    return jsonify({"ok": False, "error": message, "code": code}), status


def _require_json() -> Tuple[Dict[str, Any], Any]:
    data = request.get_json(silent=True)
    if not isinstance(data, dict):
        return {}, _json_error("Request body must be JSON object.", status=400)
    return data, None


def _validate_non_empty_text(value: Any, field: str, *, min_len: int = 1) -> Tuple[str, Any]:
    if value is None:
        return "", _json_error(f"Missing '{field}'.", status=400)

    text = str(value).strip()
    if not text:
        return "", _json_error(f"'{field}' cannot be empty.", status=400)

    # Use non-whitespace characters as the minimum length measure.
    non_ws = len("".join(text.split()))
    if non_ws < int(min_len):
        return "", _json_error(f"'{field}' is too short (min {min_len} non-whitespace chars).", status=400)

    return text, None


def _utc_timestamp_slug() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")


def _get_session_key_for_type(content_type: str) -> str:
    mapping = {
        "screenplay": "screenplay",
        "characters": "characters",
        "sound_design": "sound_design",
    }
    return mapping.get(content_type, "")


# -----------------------------
# File generation
# -----------------------------

def _make_txt_bytes(content: str) -> io.BytesIO:
    buf = io.BytesIO()
    buf.write((content or "").encode("utf-8"))
    buf.seek(0)
    return buf


def _make_pdf_bytes(title: str, content: str) -> io.BytesIO:
    # ReportLab: letter size, 11pt font, multi-page support.
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    left_margin = 54
    right_margin = 54
    top_margin = 54
    bottom_margin = 54

    font_name = "Times-Roman"
    font_size = 11
    leading = 14

    max_width = width - left_margin - right_margin

    def wrap_line(line: str) -> list[str]:
        if not line:
            return [""]
        words = line.split(" ")
        out: list[str] = []
        current = ""
        for w in words:
            candidate = (current + " " + w).strip() if current else w
            if c.stringWidth(candidate, font_name, font_size) <= max_width:
                current = candidate
            else:
                if current:
                    out.append(current)
                # If one word is extremely long, hard-split it.
                if c.stringWidth(w, font_name, font_size) > max_width:
                    chunk = ""
                    for ch in w:
                        cand = chunk + ch
                        if c.stringWidth(cand, font_name, font_size) <= max_width:
                            chunk = cand
                        else:
                            out.append(chunk)
                            chunk = ch
                    current = chunk
                else:
                    current = w
        if current:
            out.append(current)
        return out

    y = height - top_margin
    c.setFont(font_name, font_size)

    # Title
    if title:
        c.setFont("Times-Bold", 14)
        c.drawString(left_margin, y, title)
        y -= 22
        c.setFont(font_name, font_size)

    text = (content or "").replace("\r\n", "\n").replace("\r", "\n")
    for raw_line in text.split("\n"):
        for line in wrap_line(raw_line):
            if y <= bottom_margin:
                c.showPage()
                y = height - top_margin
                c.setFont(font_name, font_size)
            c.drawString(left_margin, y, line)
            y -= leading

    c.save()
    buf.seek(0)
    return buf


def _make_docx_bytes(title: str, content: str) -> io.BytesIO:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Preserve paragraphs/blank lines reasonably.
    text = (content or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")

    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# -----------------------------
# Routes
# -----------------------------

def register_routes(app: Flask) -> None:
    @app.get("/")
    def index():
        return render_template("index.html")

    @app.get("/api/health")
    def health_check():
        return jsonify({"ok": True, "status": "healthy"})

    @app.post("/set_username")
    def set_username():
        data, err = _require_json()
        if err:
            return err

        username, err = _validate_non_empty_text(data.get("username"), "username", min_len=1)
        if err:
            return err

        session["username"] = username
        return jsonify({"ok": True, "username": username})

    @app.post("/generate_content")
    def generate_content():
        data, err = _require_json()
        if err:
            return err

        # Integration mode:
        # - If a story idea is provided, run the AI pipeline via ai_engine.py
        # - Else, store already-generated outputs provided by the caller
        story_raw = data.get("story") or data.get("story_idea") or data.get("prompt")
        story = str(story_raw).strip() if story_raw is not None else ""

        if story:
            try:
                from ai_engine import AIEngineError, run_full_pipeline, run_full_pipeline_single_call  # local module
            except Exception as exc:
                return _json_error(
                    f"AI engine not available: {exc}",
                    status=500,
                    code="ai_engine_missing",
                )

            # Optional tuning knobs
            temperature_raw = data.get("temperature", None)
            min_story_chars_raw = data.get("min_story_chars", None)

            def _to_float(v: Any, default: float) -> float:
                try:
                    return float(v)
                except Exception:
                    return default

            def _to_int(v: Any, default: int) -> int:
                try:
                    return int(v)
                except Exception:
                    return default

            temperature = _to_float(temperature_raw, 0.7)
            min_story_chars = _to_int(min_story_chars_raw, int(os.getenv("MIN_STORY_CHARS", "120")))

            single_call_env = os.getenv("AI_SINGLE_CALL", "").strip().lower() in {"1", "true", "yes", "on"}
            single_call_req = bool(data.get("single_call"))
            use_single_call = single_call_env or single_call_req

            try:
                if use_single_call:
                    result = run_full_pipeline_single_call(
                        story,
                        temperature=temperature,
                        min_story_chars=min_story_chars,
                    )
                else:
                    result = run_full_pipeline(
                        story,
                        temperature=temperature,
                        min_story_chars=min_story_chars,
                    )
            except AIEngineError as exc:
                msg = str(exc)
                upper = msg.upper()
                if "RESOURCE_EXHAUSTED" in upper or "429" in upper or "RATE" in upper or "QUOTA" in upper:
                    return _json_error(msg, status=429, code="rate_limited")
                return _json_error(msg, status=400, code="ai_engine_error")
            except Exception as exc:
                return _json_error(f"AI generation failed: {exc}", status=500, code="ai_generation_failed")

            # Store outputs in session for downloads
            session["screenplay"] = result.get("screenplay") or ""
            session["characters"] = result.get("characters") or ""
            session["sound_design"] = result.get("sound_design") or ""
            session["genre_analysis"] = result.get("genre_analysis") or {}
            session["ai_meta"] = result.get("meta") or {}

            return jsonify(result)

        # No story provided: treat as a storage endpoint for already-generated outputs.
        screenplay, err = _validate_non_empty_text(data.get("screenplay"), "screenplay", min_len=1)
        if err:
            return err

        characters, err = _validate_non_empty_text(data.get("characters"), "characters", min_len=1)
        if err:
            return err

        sound_design, err = _validate_non_empty_text(data.get("sound_design"), "sound_design", min_len=1)
        if err:
            return err

        session["screenplay"] = screenplay
        session["characters"] = characters
        session["sound_design"] = sound_design

        return jsonify({"ok": True, "screenplay": screenplay, "characters": characters, "sound_design": sound_design})

    @app.post("/download/<format_type>")
    def download_file(format_type: str):
        fmt = (format_type or "").strip().lower()
        if fmt not in {"txt", "pdf", "docx"}:
            return _json_error("Invalid format_type. Must be one of: txt, pdf, docx.", status=400)

        # Frontend compatibility: allow missing JSON body and default to screenplay.
        data = request.get_json(silent=True)
        if isinstance(data, dict) and data.get("type") is not None:
            content_type, err = _validate_non_empty_text(data.get("type"), "type", min_len=1)
            if err:
                return err
            content_type = content_type.lower()
        else:
            content_type = "screenplay"

        session_key = _get_session_key_for_type(content_type)
        if not session_key:
            return _json_error(
                "Invalid download type. Must be one of: screenplay, characters, sound_design.",
                status=400,
            )

        content = session.get(session_key)
        if not content:
            return _json_error("No generated content found in session. Generate content first.", status=404)

        title_map = {
            "screenplay": "Screenplay",
            "characters": "Character Profiles",
            "sound_design": "Sound Design Plan",
        }
        title = title_map.get(content_type, "Document")

        stamp = _utc_timestamp_slug()
        filename = f"{content_type}_{stamp}.{fmt}"

        if fmt == "txt":
            buf = _make_txt_bytes(str(content))
            return send_file(
                buf,
                mimetype="text/plain; charset=utf-8",
                as_attachment=True,
                download_name=filename,
            )

        if fmt == "pdf":
            buf = _make_pdf_bytes(title, str(content))
            return send_file(
                buf,
                mimetype="application/pdf",
                as_attachment=True,
                download_name=filename,
            )

        # docx
        buf = _make_docx_bytes(title, str(content))
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )


# -----------------------------
# Error handling
# -----------------------------

def register_error_handlers(app: Flask) -> None:
    @app.errorhandler(404)
    def not_found(_):
        return _json_error("Not found.", status=404, code="not_found")

    @app.errorhandler(405)
    def method_not_allowed(_):
        return _json_error("Method not allowed.", status=405, code="method_not_allowed")

    @app.errorhandler(Exception)
    def unhandled_exception(exc: Exception):
        # Avoid leaking stack traces in production.
        msg = str(exc) if app.config.get("DEBUG") else "Internal server error."
        return _json_error(msg, status=500, code="internal_error")


app = create_app()


if __name__ == "__main__":
    host = os.getenv("FLASK_HOST", "127.0.0.1")
    port = int(os.getenv("FLASK_PORT", "5000"))
    app.run(host=host, port=port, debug=app.config.get("DEBUG", False))
